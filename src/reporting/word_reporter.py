"""Word-Report-Generator (.docx)."""

from datetime import datetime
from typing import List

from src.models.testcase import TestCaseResult


def generate_word_report(
    results: List[TestCaseResult],
    excel_file: str,
    output_path: str,
) -> str:
    """Erstellt einen Word-Report (.docx)."""
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # Titel
    doc.add_heading("Testlauf-Zusammenfassung", level=0)

    # Metadaten
    doc.add_paragraph(f"Datum/Uhrzeit: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Input-Datei: {excel_file}")
    doc.add_paragraph(f"Testfälle gesamt: {len(results)}")

    pass_count = sum(1 for r in results if r.overall_pass)
    fail_count = len(results) - pass_count
    doc.add_paragraph(f"Pass: {pass_count} | Fail: {fail_count}")

    doc.add_heading("Ergebnisse pro Testfall", level=1)

    # Tabelle
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    headers = [
        "TestcaseID", "Titel", "Zahlungstyp",
        "XSD-Status", "Business Rules", "Ergebnis", "Bemerkungen",
    ]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for r in results:
        row = table.add_row()
        row.cells[0].text = r.testcase_id
        row.cells[1].text = r.titel
        row.cells[2].text = r.payment_type.value
        row.cells[3].text = "OK" if r.xsd_valid else "FEHLER"

        # Business Rule Status
        failed_rules = [br for br in r.business_rule_results if not br.passed]
        if failed_rules:
            rule_texts = [f"{br.rule_id}: {br.rule_description}" for br in failed_rules]
            row.cells[4].text = "\n".join(rule_texts)
        else:
            row.cells[4].text = "Alle bestanden"

        row.cells[5].text = "Pass" if r.overall_pass else "Fail"
        row.cells[6].text = r.remarks or ""

    report_path = f"{output_path}/Testlauf_Zusammenfassung.docx"
    doc.save(report_path)
    return report_path


def generate_txt_report(
    results: List[TestCaseResult],
    excel_file: str,
    output_path: str,
) -> str:
    """Erstellt einen Text-Report als Fallback."""
    lines = [
        "=" * 70,
        "TESTLAUF-ZUSAMMENFASSUNG",
        "=" * 70,
        f"Datum/Uhrzeit: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}",
        f"Input-Datei: {excel_file}",
        f"Testfälle gesamt: {len(results)}",
        f"Pass: {sum(1 for r in results if r.overall_pass)}",
        f"Fail: {sum(1 for r in results if not r.overall_pass)}",
        "",
        "-" * 70,
    ]

    for r in results:
        lines.append(f"TestcaseID: {r.testcase_id}")
        lines.append(f"  Titel: {r.titel}")
        lines.append(f"  Zahlungstyp: {r.payment_type.value}")
        lines.append(f"  XSD-Status: {'OK' if r.xsd_valid else 'FEHLER'}")
        if r.xsd_errors:
            for e in r.xsd_errors:
                lines.append(f"    XSD-Fehler: {e}")

        failed_rules = [br for br in r.business_rule_results if not br.passed]
        if failed_rules:
            lines.append("  Business Rules (fehlgeschlagen):")
            for br in failed_rules:
                lines.append(f"    {br.rule_id}: {br.rule_description}")
                if br.details:
                    lines.append(f"      Details: {br.details}")
        else:
            lines.append("  Business Rules: Alle bestanden")

        lines.append(f"  Ergebnis: {'Pass' if r.overall_pass else 'Fail'}")
        if r.remarks:
            lines.append(f"  Bemerkungen: {r.remarks}")
        lines.append("-" * 70)

    report_path = f"{output_path}/Testlauf_Zusammenfassung.txt"
    with open(report_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))

    return report_path
